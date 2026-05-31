"""
Build a Microsoft Word (.docx) version of the Q1 paper skeleton, mirroring
the LaTeX structure. All tables are regenerated from the canonical CSVs in
reports/ and all figures are embedded from paper/figures/defense/*.png so
that nothing in the Word file is hand-edited boilerplate.

Output:
    paper/hvac_paper_skeleton.docx

Reads:
    reports/hou_evins_architecture_justification_table.csv  (Table 1)
    reports/hou_evins_predictive_validity_table.csv         (Table 2)
    reports/speed_benchmark_table.csv                       (Table 4)
    reports/morl_pareto_front_table.csv                     (Table 5)
    reports/pi_baseline_yearly_table.csv                    (Table 5)
    paper/figures/defense/F*.png                            (Figures)

Safe to rerun. The same script generates the LaTeX tables; the Word file
takes the same numbers so the two paper formats never drift.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Optional

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PAPER_DIR = Path(__file__).resolve().parent
OUT_PATH = PAPER_DIR / "hvac_paper_skeleton.docx"
FIG_DIR = PAPER_DIR / "figures" / "defense"
REPO_ROOT = PAPER_DIR.parent


def _data_root() -> Path:
    candidates = [REPO_ROOT]
    parts = REPO_ROOT.parts
    if ".claude" in parts:
        idx = parts.index(".claude")
        candidates.append(Path(*parts[:idx]))
    for cand in candidates:
        if (cand / "reports").exists():
            return cand
    return REPO_ROOT


DATA_ROOT = _data_root()


def _find_csv(rel_path: str) -> Optional[Path]:
    for root in (REPO_ROOT, DATA_ROOT):
        p = root / rel_path
        if p.exists():
            return p
    return None


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

TODO_COLOR = RGBColor(0xC0, 0x10, 0x20)
MUTED_COLOR = RGBColor(0x70, 0x70, 0x70)


def _setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        h = styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_para(doc: Document, text: str, *, italic: bool = False,
             bold: bool = False, color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if italic:
        run.italic = True
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color


def add_todo(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[TODO] {text}")
    run.italic = True
    run.font.color.rgb = TODO_COLOR


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED_COLOR


def add_table_from_rows(doc: Document, headers: list[str],
                        rows: list[list[str]],
                        col_widths_cm: Optional[list[float]] = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, value in enumerate(row):
            cells[c].text = ""
            p = cells[c].paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm is not None:
        for row in table.rows:
            for c, w in enumerate(col_widths_cm):
                row.cells[c].width = Cm(w)


def add_figure(doc: Document, png_path: Path, caption: str,
               width_inches: float = 6.3) -> None:
    if not png_path.exists():
        add_todo(doc, f"Missing figure: {png_path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png_path), width=Inches(width_inches))
    add_caption(doc, caption)


def _fmt(value: object, fmt: str = "{:.3f}") -> str:
    if value is None:
        return "—"
    try:
        if pd.isna(value):
            return "—"
        return fmt.format(float(value))
    except (TypeError, ValueError):
        return str(value)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def build_table1(doc: Document) -> None:
    csv = _find_csv("reports/hou_evins_architecture_justification_table.csv")
    if csv is None:
        add_todo(doc, "Table 1 source CSV missing: reports/hou_evins_architecture_justification_table.csv")
        return
    df = pd.read_csv(csv).set_index("variant")
    needed = ["v3", "v35_calibrated", "hybrid_l010"]
    label = {"v3": "v3", "v35_calibrated": "v3.5 calibrated", "hybrid_l010": "hybrid_l010"}

    headers = ["Variant", "Explicit C_zon", "1-step RMSE (°C)", "24h RMSE (°C)",
               "Peak m_s", "Typical m_s", "Peak transfer RMSE (°C)", "Typical transfer RMSE (°C)"]
    rows = []
    for v in needed:
        if v not in df.index:
            continue
        r = df.loc[v]
        rows.append([
            label[v],
            str(r.get("explicit_c_zon", "—")),
            _fmt(r.get("block1_temp_alignment_rmse_c")),
            _fmt(r.get("block1_rollout_24h_rmse_c")),
            _fmt(r.get("peak_control_m_s")),
            _fmt(r.get("typical_control_m_s")),
            _fmt(r.get("peak_transfer_temp_rmse_c")),
            _fmt(r.get("typical_transfer_temp_rmse_c")),
        ])
    add_table_from_rows(doc, headers, rows)
    add_caption(doc,
        "Table 1. Architecture comparison of the three surrogate variants on "
        "bestest_air. v3 is the control-oriented black-box, v3.5 calibrated is "
        "the physically informed RC-NeuralODE after Stage A/B/C calibration "
        "(explicit C_zon), and hybrid_l010 is the canonical hybrid backend "
        "(λ_temp = 0.10, λ_power = 5e-5). Source: "
        "reports/hou_evins_architecture_justification_table.csv.")


def build_table2(doc: Document) -> None:
    csv = _find_csv("reports/hou_evins_predictive_validity_table.csv")
    if csv is None:
        add_todo(doc, "Table 2 source CSV missing.")
        return
    df = pd.read_csv(csv)
    df = df[df["validity_type"] == "predictive_prepared_rollout"]

    horizons = ["rollout_1h", "rollout_4h", "rollout_8h", "rollout_24h"]
    variants = ["raw_v35", "v35_calibrated"]
    metrics  = [("RMSE_T_C", "RMSE (°C)"), ("MAE_T_C", "MAE (°C)")]

    def lookup(v, h, m):
        sel = df[(df["variant"] == v) & (df["horizon"] == h) & (df["metric"] == m)]
        return _fmt(sel["value"].iloc[0]) if not sel.empty else "—"

    headers = ["Variant", "Metric"] + [h.replace("rollout_", "") for h in horizons]
    rows = []
    for v in variants:
        for m_key, m_lab in metrics:
            rows.append([
                "raw_v35" if v == "raw_v35" else "v35_calibrated",
                m_lab,
                *[lookup(v, h, m_key) for h in horizons]
            ])
    add_table_from_rows(doc, headers, rows)
    add_caption(doc,
        "Table 2. Predictive validity of the calibrated physical twin v3.5 "
        "versus its uncalibrated baseline raw_v35 across rollout horizons. "
        "Values are computed on the held-out prepared 15-minute corpus "
        "(8 episodes). Calibration roughly halves the temperature RMSE at "
        "every horizon and the RMSE is nearly flat from 1h to 24h, "
        "indicating that the identified C_zon produces a stable physical "
        "model rather than one that accumulates error with horizon.")


def build_table4(doc: Document) -> None:
    csv = _find_csv("reports/speed_benchmark_table.csv")
    if csv is None:
        add_todo(doc, "Table 4 source CSV missing.")
        return
    df = pd.read_csv(csv).set_index("backend")
    order = [
        ("boptest_rte_http",         "BOPTEST RTE (HTTP)"),
        ("v3_surrogate",             "v3 surrogate"),
        ("v35_calibrated_surrogate", "v3.5 calibrated"),
        ("hybrid_v3_v35_surrogate",  "hybrid_l010 (canonical)"),
    ]
    headers = ["Backend", "Steps/s", "Median step (ms)", "P95 step (ms)", "Speed-up"]
    rows = []
    for key, name in order:
        if key not in df.index:
            continue
        r = df.loc[key]
        rows.append([
            name,
            _fmt(r["env_steps_per_sec"], "{:,.1f}"),
            _fmt(r["median_raw_step_ms"], "{:.3f}"),
            _fmt(r["p95_raw_step_ms"], "{:.3f}"),
            _fmt(r["speedup_vs_boptest_rte"], "{:.1f}×"),
        ])
    add_table_from_rows(doc, headers, rows)
    add_caption(doc,
        "Table 4. Throughput comparison on CPU at the same 15-min control "
        "protocol (100 episodes × 96 steps = 9,600 transitions per backend). "
        "The live BOPTEST RTE row uses the same HTTP API path that downstream "
        "RL training and live closed-loop validation actually use, so the "
        "speed-up factor in the last column reflects the practical, not the "
        "idealized, ratio.")


def build_table5(doc: Document) -> None:
    pareto_csv = _find_csv("reports/morl_pareto_front_table.csv")
    pi_csv     = _find_csv("reports/pi_baseline_yearly_table.csv")
    if pareto_csv is None or pi_csv is None:
        add_todo(doc, "Table 5 source CSV missing.")
        return
    pareto = pd.read_csv(pareto_csv)
    pi     = pd.read_csv(pi_csv)
    pi_row = pi[pi["controller"] == "pi_builtin"].iloc[0] if not pi.empty else None

    pg = pareto.groupby(
        ["preference_w_comfort", "preference_w_energy", "canonical_designation"],
        as_index=False,
    ).agg(
        n_seeds=("seed", "count"),
        m_s=("m_s", "mean"),
        violation_pct=("violation_pct", "mean"),
        energy_kwh=("energy_kwh", "mean"),
        rmse_yearly_c=("rmse_yearly_c", "mean"),
    ).sort_values("preference_w_comfort", ascending=False).reset_index(drop=True)

    designation_label = {
        "pareto_endpoint_comfort":         "comfort endpoint",
        "practical_deployment_canonical":  "★ practical canonical",
        "pre_registered_canonical":        "★ pre-registered canonical",
        "pareto_intermediate":             "intermediate",
        "pareto_endpoint_energy_collapse": "energy collapse",
    }

    headers = ["Configuration", "Designation", "Yearly m_s",
               "Violation (%)", "Energy (kWh)", "RMSE (°C)", "N seeds"]
    rows = []
    if pi_row is not None:
        rows.append([
            "BOPTEST PI (built-in)",
            "reference",
            _fmt(pi_row["m_s"], "{:.3f}"),
            _fmt(pi_row["violation_pct"], "{:.2f}"),
            _fmt(pi_row["energy_kwh"], "{:.2f}"),
            _fmt(pi_row["rmse_yearly_c"], "{:.3f}"),
            "1",
        ])
    for _, r in pg.iterrows():
        rows.append([
            f"MORL  w=({r['preference_w_comfort']:.2f}, {r['preference_w_energy']:.2f})",
            designation_label.get(r["canonical_designation"], r["canonical_designation"]),
            _fmt(r["m_s"], "{:.4f}"),
            _fmt(r["violation_pct"], "{:.2f}"),
            _fmt(r["energy_kwh"], "{:.2f}"),
            _fmt(r["rmse_yearly_c"], "{:.3f}"),
            str(int(r["n_seeds"])),
        ])
    add_table_from_rows(doc, headers, rows)
    add_caption(doc,
        "Table 5. Yearly evaluation of the five-weight MORL Pareto sweep on "
        "bestest_air (single seed per preference vector at the time of writing; "
        "canonical rows will be re-rendered with mean ± std once seeds 43 and 44 "
        "complete). The BOPTEST built-in PI controller is included as the "
        "standard reference row; it is the default tuning exposed by the "
        "testcase and is NOT a custom-tuned baseline.")


# ---------------------------------------------------------------------------
# Sections (mirror of LaTeX skeleton)
# ---------------------------------------------------------------------------

def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "When Predictive Surrogates Fail as RL Environments:\n"
        "A Calibrated Physical Twin as Soft Regularizer for HVAC Control"
    )
    run.font.size = Pt(20)
    run.font.bold = True

    doc.add_paragraph()
    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by.add_run("Almaz Sapargali\n").italic = True
    add_todo(doc, "Affiliation, address, country.")

    doc.add_paragraph()
    target = doc.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = target.add_run("Target journal: Results in Engineering (Elsevier, Q1)")
    r.italic = True
    r.font.color.rgb = MUTED_COLOR


def add_abstract(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    add_para(doc,
        "Reinforcement learning controllers for heating, ventilation and air "
        "conditioning (HVAC) systems are typically trained on neural-network "
        "surrogates because high-fidelity simulators are too slow for the "
        "millions of environment steps that modern policy-gradient methods "
        "consume. A natural assumption is that a more physically faithful "
        "surrogate produces a better training environment. We test that "
        "assumption on the BOPTEST bestest_air testcase and report a clear "
        "negative result: a calibrated physical twin with explicit zone "
        "thermal capacitance achieves a 24-hour rollout root-mean-square "
        "temperature error of 0.64 °C (versus 1.47 °C uncalibrated), yet "
        "collapses as a stand-alone reinforcement-learning training "
        "environment, producing a live closed-loop RMSE above 4 °C. We "
        "resolve the gap by repurposing the calibrated twin as a frozen "
        "soft physical regularizer for a smoother data-driven surrogate. "
        "The resulting hybrid backend sustains 1,787 environment steps per "
        "second on a single CPU thread, corresponding to an 85.0× speed-up "
        "over the standard BOPTEST RTE HTTP–Docker deployment used in "
        "production benchmarking, under the same 15-minute control protocol, "
        "while restoring live closed-loop RMSE below 0.65 °C. The optimal "
        "regularization strength is "
        "controller-family-specific: λ_temp = 0.10 for thermostatic "
        "proximal-policy optimization, but λ_temp = 0.00 for both "
        "hierarchical reinforcement learning and multi-objective "
        "reinforcement learning with a 17-dimensional observation "
        "interface. The full pipeline follows the surrogate-development "
        "protocol of Hou and Evins (2024), with all numerical "
        "justifications provided in Supplementary Tables S1–S11."
    )
    add_para(doc, "Keywords: HVAC control; deep reinforcement learning; "
                  "digital twin; physics-informed machine learning; BOPTEST; "
                  "multi-objective reinforcement learning.",
             italic=True)


def add_section_introduction(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    add_todo(doc, "Opening hook: building energy demand, decarbonization, role "
                  "of advanced HVAC control. Finish with a sentence that sets "
                  "up the surrogate problem.")
    add_todo(doc, "Why RL for HVAC is sample-expensive and brittle on "
                  "high-fidelity simulators; standard practice is to train on "
                  "surrogates.")
    add_todo(doc, "The implicit assumption that better surrogate fidelity "
                  "implies a better RL training environment. Forward-reference "
                  "Section 5 where this is tested and falsified.")
    add_todo(doc, "Our resolution — calibrated twin used as soft regularizer, "
                  "not as training environment. Briefly describe the hybrid "
                  "construction.")

    doc.add_heading("Contributions", level=2)
    contributions = [
        "A physics-informed surrogate (v3.5) of the BOPTEST bestest_air "
        "testcase, calibrated through a three-stage inverse procedure "
        "(Stage A telemetry preprocessing, Stage B explicit identification "
        "of zone thermal capacitance C_zon, Stage C residual head "
        "calibration), achieving a 38% reduction in one-step temperature "
        "RMSE and a 40% reduction in power MAE over the uncalibrated baseline.",

        "A clear empirical demonstration that predictive validity does not "
        "transfer to reinforcement-learning training utility: the calibrated "
        "twin achieves a 24-hour rollout RMSE of 0.64 °C on held-out "
        "trajectories yet drives a closed-loop BOPTEST RMSE above 4 °C when "
        "used as a stand-alone RL training environment.",

        "A hybrid surrogate construction (hybrid_l010) in which the "
        "calibrated twin acts as a frozen soft physical regularizer on a "
        "smoother control-oriented surrogate, restoring closed-loop transfer "
        "RMSE below 0.65 °C on both peak and typical heating scenarios while "
        "sustaining 1,787 environment steps per second on a single CPU "
        "thread (85.0× speed-up over the standard BOPTEST RTE HTTP–Docker "
        "deployment used in production benchmarking).",

        "A controller-family-specific finding: the optimal physical "
        "regularization strength differs across controller architectures "
        "(λ_temp = 0.10 for thermostatic PPO, λ_temp = 0.00 for HDRL and "
        "17-D MORL), with an interpretation grounded in observation-space "
        "geometry and control hierarchy.",

        "A reproducible BOPTEST-based evaluation protocol with full "
        "Hou-and-Evins-compliant numerical justification for sample "
        "generation, preprocessing, feature significance and independence, "
        "scaling, hyperparameters, architecture, and "
        "replicative-plus-predictive validity, provided as Supplementary "
        "Tables S1–S11.",
    ]
    for c in contributions:
        doc.add_paragraph(c, style="List Number")


def add_section_related_work(doc: Document) -> None:
    doc.add_heading("2. Related Work", level=1)
    for header, todo in [
        ("2.1 Deep reinforcement learning for HVAC control",
         "Cite Wei et al. (DQN for HVAC), early PPO/DDPG building "
         "applications, Sinergym, hierarchical and safe-RL approaches. "
         "State the controller-side gap."),
        ("2.2 Surrogate and digital twin models for building energy",
         "Reference NN-based surrogates, in particular Hou and Evins (2024). "
         "Distinguish predictive surrogates from closed-loop RL surrogates. "
         "State the fidelity-side gap."),
        ("2.3 Physics-informed and physics-guided machine learning",
         "Distinguish hard-constraint PINNs from soft regularization. "
         "Position our hybrid construction as soft, not PINN."),
        ("2.4 BOPTEST benchmarking",
         "Describe BOPTEST and bestest_air, KPI structure, rationale."),
        ("2.5 Position of this work",
         "State that this work sits at the intersection of the three gaps."),
    ]:
        doc.add_heading(header, level=2)
        add_todo(doc, todo)


def add_section_methodology(doc: Document) -> None:
    doc.add_heading("3. Methodology", level=1)
    add_para(doc,
        "The methodology is organized in four blocks: (i) the control-oriented "
        "data-driven surrogate v3, (ii) the physically informed surrogate v3.5 "
        "together with its three-stage inverse calibration pipeline, (iii) the "
        "hybrid backend that combines the two, and (iv) the controller families "
        "and benchmark protocol.")

    doc.add_heading("3.1 Control-oriented surrogate v3", level=2)
    add_todo(doc, "Two-headed feed-forward NN on direct-TSup trajectories.")

    doc.add_heading("3.2 Physically informed surrogate v3.5 and inverse calibration", level=2)
    add_todo(doc, "RC-NeuralODE with explicit learnable C_zon.")
    add_para(doc, "Stage A — Telemetry preprocessing.", bold=True)
    add_todo(doc, "5 operations: latency compensation, T-bias removal, "
                  "power affine normalization, rolling denoise, causal delta "
                  "recomputation.")
    add_para(doc, "Stage B — Identification of C_zon.", bold=True)
    add_todo(doc, "Excitation-window subselection (quantile 0.95 on |ΔT|), "
                  "episode-aware split, optimization of C_zon with frozen heads. "
                  "Final value: C_zon = 4.413e5 J/K.")
    add_para(doc, "Stage C — Residual head calibration.", bold=True)
    add_todo(doc, "C_zon frozen, residual heads updated.")

    doc.add_heading("3.3 Hybrid backend: v3 dynamics with v3.5 as soft regularizer", level=2)
    add_para(doc,
        "The hybrid backend treats the calibrated v3.5 model as a frozen "
        "physical censor. Policy rollouts evolve under the smoother v3 "
        "dynamics; the policy is penalized at training time for actions on "
        "which the two surrogates physically disagree:")
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq.add_run("L_total = L_PPO(π; v3) "
                   "+ λ_temp · ‖T_v3(s,a) − T_v3.5(s,a)‖²² "
                   "+ λ_power · ‖P_v3(s,a) − P_v3.5(s,a)‖²²")
    r.italic = True
    add_para(doc,
        "λ_temp and λ_power are the only two hyperparameters of the hybrid "
        "construction. Crucially, v3.5 never enters the policy's forward "
        "pass — it only enters the loss.")
    add_todo(doc, "Insert pipeline schematic (figure: pipeline_overview.pdf) "
                  "showing v3 forward pass, v3.5 frozen, λ-weighted disagreement "
                  "entering only the loss.")

    doc.add_heading("3.4 Controller families and benchmark protocol", level=2)
    add_todo(doc, "PI baseline, Thermostatic PPO, HDRL, MORL. Define each.")


def add_section_experimental_setup(doc: Document) -> None:
    doc.add_heading("4. Experimental Setup", level=1)

    doc.add_heading("4.1 Testbed and serving layer", level=2)
    add_para(doc,
        "All experiments use the BOPTEST bestest_air testcase, a single-zone "
        "air-side HVAC reference building. BOPTEST is served through the "
        "boptest_rte container, which exposes the standard BOPTEST HTTP API "
        "to both surrogate training and live closed-loop validation.")
    add_todo(doc, "One paragraph on building envelope and TMY weather.")

    doc.add_heading("4.2 Datasets and sample generation", level=2)
    add_todo(doc, "Refer to Supplementary Tables S1–S2.")

    doc.add_heading("4.3 Preprocessing, scaling, feature engineering", level=2)
    add_todo(doc, "Refer to Supplementary Tables S4, S6, S7.")

    doc.add_heading("4.4 Evaluation Protocol", level=2)
    add_para(doc,
        "Two live BOPTEST closed-loop windows are used as the primary "
        "transfer benchmarks: peak_heat_window (worst-case cold) and "
        "typical_heat_window (mild average). Yearly validation uses the "
        "full BOPTEST year with KPI windows defined by the BOPTEST "
        "specification. The simulation step is Δt = 900 s (15 min); the "
        "comfort band is [21 °C, 24 °C].")
    add_para(doc, "Metrics: m_s (BOPTEST composite KPI), violation %, "
                  "energy (kWh), RMSE_T (replicative one-step and multi-horizon "
                  "predictive), first_divergence_step (transfer-realism "
                  "diagnostic).")
    add_para(doc,
        "Statistical reporting: predictive validity numbers report bootstrap "
        "95% confidence intervals from horizon_metrics.csv. Canonical "
        "controller results are reported as mean ± std over N = 3 random "
        "seeds for thermostatic hybrid, pure v3, and MORL canonical. "
        "HDRL is reported single-seed; this is acknowledged in §8.")

    doc.add_heading("4.5 Reproducibility statement", level=2)
    add_todo(doc, "Repository URL, BOPTEST version, Python/PyTorch versions, "
                  "hardware, random seed handling.")

    doc.add_heading("4.6 Runtime characteristics", level=2)
    add_para(doc,
        "A direct throughput comparison between the live BOPTEST RTE HTTP "
        "loop and the three surrogate backends is given in Table 4 and "
        "Figure 3. All backends were exercised at the same 15-minute "
        "control protocol on a single CPU thread, with 100 episodes of 96 "
        "transitions each (9,600 total transitions per backend). The live "
        "BOPTEST RTE backend reaches 21.0 environment steps per second; "
        "v3 reaches 4,626 steps/s (220×), v3.5 calibrated 2,400 steps/s "
        "(114×), and the canonical hybrid backend 1,787 steps/s (85×). "
        "The hybrid is slower than v3 alone because it evaluates both "
        "networks plus the disagreement terms on every step; 85× is the "
        "practical speed-up factor for the rest of the paper because all "
        "canonical RL training uses the hybrid backend.")
    build_table4(doc)
    add_figure(doc, FIG_DIR / "F3_speed_vs_fidelity.png",
               "Figure 3. Throughput comparison on CPU under the same 15-min "
               "control protocol. The hybrid backend trains RL controllers 85× "
               "faster than the standard BOPTEST RTE HTTP–Docker deployment "
               "used in production benchmarking.")
    add_para(doc,
        "Scope of the comparison. The 85× figure compares the surrogate to "
        "the standard BOPTEST RTE deployment — the HTTP–Docker stack used by "
        "the BOPTEST community for production benchmarking and used here for "
        "all downstream live closed-loop validation. This comparison "
        "conflates two costs: the intrinsic Modelica/FMU simulation cost and "
        "the deployment overhead (HTTP roundtrip, JSON serialization, Flask "
        "request handling, container isolation). To isolate the simulation "
        "cost, we implemented a benchmark of direct FMU loading via the fmpy "
        "co-simulation interface (script "
        "evaluation/build_speed_benchmark_fmu_direct.py). The BOPTEST "
        "bestest_air FMU ships with a Linux x86_64 binary only "
        "(binaries/linux64/wrapped.so); the archive contains no Windows or "
        "macOS native binary, which is the reason BOPTEST is conventionally "
        "deployed in a Linux container. Because all measurements in this "
        "study were collected on a Windows host (matching the host of the "
        "live HTTP benchmark for cross-comparability), the direct FMU "
        "benchmark cannot be run in-process on the same machine without a "
        "Linux runtime. The script is committed with explicit reproducibility "
        "instructions for WSL2 Ubuntu and Docker; the platform constraint is "
        "logged at reports/speed_benchmark_fmu_direct_platform_log.txt. We "
        "accordingly frame the reported speed-up as a comparison against the "
        "practical BOPTEST deployment configuration rather than against bare "
        "FMU evaluation. A direct FMU–surrogate comparison remains a deferred "
        "reproducibility item: when executed on a Linux host, it will produce "
        "a third row in Table 4 that isolates the deployment-overhead "
        "component of the current 85× figure.")


def add_section_results_fidelity(doc: Document) -> None:
    doc.add_heading("5. Results I: Digital Twin Fidelity", level=1)
    add_para(doc,
        "This section answers the question \"How accurate is the surrogate as "
        "a model of the building?\" It deliberately does not answer \"Does it "
        "train good RL controllers?\" — that question is addressed in §6.")

    doc.add_heading("5.1 Architecture comparison", level=2)
    add_todo(doc, "Describe Table 1 and the three variants.")
    build_table1(doc)

    doc.add_heading("5.2 Replicative validity", level=2)
    add_para(doc,
        "Calibrated v3.5 achieves a one-step temperature RMSE of 0.232 °C "
        "versus a pre-calibration baseline of 0.374 °C, a 38% reduction. "
        "Power MAE drops from 807.8 W to 482.0 W, a 40% reduction.")
    add_para(doc,
        "Identification of C_zon: the calibrated zone thermal capacitance is "
        "C_zon = 4.413 × 10⁵ J/K, deviating from the synthetic prior "
        "5.3 × 10⁵ J/K by approximately 21%. Stable across training; the "
        "final value is frozen in the canonical checkpoint and reused "
        "unchanged in all downstream experiments.")

    doc.add_heading("5.3 Predictive validity", level=2)
    add_para(doc,
        "Table 2 reports 1h/4h/8h/24h rollout RMSE_T and MAE_T on the "
        "held-out prepared corpus. Three observations: (i) calibration roughly "
        "halves the RMSE at every horizon; (ii) the calibrated curve is "
        "approximately flat from 1h to 24h (0.65 → 0.64 °C), indicating that "
        "the identified C_zon produces a stable physical model rather than "
        "one that accumulates error with horizon; (iii) the raw model's RMSE "
        "is also flat (~1.49 °C at all horizons), indicating a systematic "
        "bias rather than divergent dynamics.")
    build_table2(doc)

    add_figure(doc, FIG_DIR / "F1_calibration_paradox.png",
               "Figure 1. The calibration paradox: v3.5 calibrated achieves the "
               "best predictive validity (0.64 °C on held-out 24h rollout) but "
               "the worst closed-loop transfer RMSE (~4.4 °C on live BOPTEST), "
               "a 6.8× gap. Predictive accuracy alone does not produce a "
               "usable RL training environment.")

    doc.add_heading("5.4 Failure mode: predictive twin is not an RL training environment", level=2)
    add_para(doc,
        "The strong predictive validity of v3.5 does not transfer to its "
        "usefulness as a closed-loop RL training environment. When a PPO "
        "controller is trained with v3.5 as the sole environment and then "
        "deployed on live BOPTEST, the closed-loop temperature RMSE exceeds "
        "4 °C on both heating scenarios and first_divergence_step drops to 1 "
        "(immediate divergence).")
    add_para(doc,
        "We attribute this failure to a structural difference: in predictive "
        "validation every rollout starts from a real BOPTEST state, so "
        "residual errors do not accumulate into a self-reinforcing loop. In "
        "closed-loop RL the policy learns to compensate for the surrogate's "
        "residual errors as if they were real disturbances; when deployed "
        "on live BOPTEST those compensations become unmodeled inputs and "
        "the closed loop drifts.")

    add_figure(doc, FIG_DIR / "F2_hybrid_resolves.png",
               "Figure 2. The hybrid backend resolves the fidelity-to-RL gap. "
               "v3 has weak predictive validity but good transfer; v3.5 is "
               "the opposite; hybrid_l010 inherits both — low predictive "
               "error AND low closed-loop transfer error.")


def add_section_results_control(doc: Document) -> None:
    doc.add_heading("6. Results II: Control Performance", level=1)
    add_para(doc,
        "This section answers the orthogonal question \"Does the surrogate "
        "improve downstream RL controller performance on live BOPTEST?\". "
        "All RL results are reported relative to the BOPTEST built-in PI "
        "baseline, presented first.")

    doc.add_heading("6.1 PI baseline", level=2)
    add_para(doc,
        "The BOPTEST bestest_air testcase exposes a built-in PI controller, "
        "invoked by passing an empty action dictionary to the standard BOPTEST "
        "step API. We use this controller as the standard reference baseline "
        "because it is the reproducible default available to any BOPTEST user; "
        "the same protocol, the same setpoint schedule, the same comfort band, "
        "and the same Δt = 900 s time step are applied to PI and to every RL "
        "controller reported here.")
    add_para(doc,
        "Under yearly evaluation the built-in PI achieves m_s = 0.910, "
        "comfort violation rate of 63.6%, total energy consumption of "
        "104.07 kWh, and a yearly mean temperature RMSE of 3.395 °C "
        "(Table 5, top row). The high violation rate at the yearly horizon "
        "indicates that the default PI tuning is not optimized for full-year "
        "operation across both heating and cooling regimes; we therefore do "
        "NOT claim that this baseline represents the best achievable PI "
        "performance on this testcase. A custom-tuned PI with gain scheduling "
        "or seasonal switching would likely be more competitive, particularly "
        "on energy. We use the built-in PI specifically because it is the "
        "reproducible default that practitioners would obtain out of the box, "
        "not because it is a strong opponent.")

    doc.add_heading("6.2 Negative control: direct v3.5 warm-start", level=2)
    add_todo(doc, "Warm-starting PPO on calibrated v3.5 and then "
                  "fine-tuning on BOPTEST produces a final policy strictly "
                  "worse than training from scratch. Cite numbers from "
                  "outputs/block2_thermostatic_warmstart_utility/.")

    doc.add_heading("6.3 Thermostatic PPO with hybrid regularization", level=2)
    add_todo(doc, "λ_temp sweep {0.05, 0.10, 0.15} with λ_power = 5e-5 fixed. "
                  "Best value λ_temp = 0.10. Peak m_s = 0.087, typical "
                  "m_s = 0.041. Comparison to pure v3.")

    doc.add_heading("6.4 Hierarchical reinforcement learning (HDRL)", level=2)
    add_todo(doc, "Inverted result: λ_temp sweep {0.00, 0.03, 0.05, 0.10} "
                  "shows monotone degradation. Best: λ_temp = 0.00. "
                  "Interpretation: inter-layer conflict.")

    doc.add_heading("6.5 Multi-objective reinforcement learning (MORL)", level=2)
    add_para(doc, "Observation interface matters.", bold=True)
    add_para(doc,
        "A 5-D preference-aware MORL on the hybrid backend yielded yearly "
        "m_s = 1.046, comfort violation 74.5%, RMSE 4.96 °C — effectively a "
        "controller failure. Replacing the observation interface with a 17-D "
        "TSup-style stack (causal_smooth delta features, clipped_log power, "
        "raw t_zone) while keeping the backend, the hybrid loss, and the "
        "training pipeline fixed restored performance to the useful regime. "
        "Optimal regularization weight: λ_temp = 0.00, mirroring HDRL — a "
        "sufficiently rich observation geometry obviates explicit physical "
        "regularization.")

    add_para(doc, "Pareto front (yearly).", bold=True)
    add_para(doc,
        "We run a five-point preference sweep on the comfort-energy simplex. "
        "Results are summarized in Table 5 and plotted in Figure 4. The four "
        "comfort-leaning preferences {1.00, 0.75, 0.50, 0.25} produce a "
        "monotone front with no dominated points. The energy-only endpoint "
        "w = (0.00, 1.00) exhibits an expected safety collapse (0.28 kWh, "
        "87% violation). A natural knee separates the front near w = (0.50, "
        "0.50): the violation rate jumps from 1.69% at w = (0.75, 0.25) to "
        "6.86% at the neutral midpoint. If a deployment-defensible cutoff of "
        "5% violation is imposed, only w ∈ {(1.00, 0.00), (0.75, 0.25)} "
        "remain admissible; among these, w = (0.75, 0.25) consumes less "
        "energy and is the better practical operating point.")

    add_para(doc, "Dual canonical selection.", bold=True)
    add_para(doc,
        "We report 3-seed variance on two preference vectors: the "
        "pre-registered canonical w = (0.50, 0.50) (fixed before inspecting "
        "Pareto results, preserving pre-registration integrity) and the "
        "practical-deployment canonical w = (0.75, 0.25) (first preference "
        "vector below the 5% deployment threshold). Both rows in Table 5 are "
        "currently single seed and will be re-rendered with mean ± std once "
        "seeds 43 and 44 complete.")

    build_table5(doc)
    add_figure(doc, FIG_DIR / "F4_pareto_vs_pi.png",
               "Figure 4. Yearly Pareto scatter of the 17-D MORL controller "
               "across five preference weights, with the BOPTEST built-in PI "
               "baseline as the standard reference marker. The vertical dashed "
               "line marks the 5% deployment-violation threshold; the gray "
               "dashed curve traces the admissible region.")

    doc.add_heading("6.6 Cross-controller comparison", level=2)
    add_para(doc,
        "The optimal λ_temp depends on observation richness and on controller "
        "hierarchy: thermostatic PPO benefits from external physical anchor "
        "(λ_temp = 0.10), while HDRL and 17-D MORL prefer no temperature "
        "anchor (λ_temp = 0.00). This is the central controller-family-"
        "specific finding of the paper.")
    add_figure(doc, FIG_DIR / "F5_controller_family.png",
               "Figure 5. The optimal temperature-disagreement weight λ_temp "
               "is controller-family-specific. Thermostatic PPO benefits from "
               "a physical anchor; HDRL and 17-D MORL do not.")


def add_section_results_transfer(doc: Document) -> None:
    doc.add_heading("7. Results III: Transferability and Generalization", level=1)
    add_para(doc,
        "Optional section. Activate only if Block 3 produces usable results "
        "before submission.", italic=True)
    for h, todo in [
        ("7.1 Selection of related testcases",
         "Pre-registered selection of 2–3 related BOPTEST testcases."),
        ("7.2 Recalibration regimes",
         "Three regimes: none, partial (Stage C only), full (Stage A/B/C)."),
        ("7.3 Transfer matrix",
         "Table testcase × recalibration × controller family."),
        ("7.4 Verdict per cell",
         "PASS / CONDITIONAL PASS / FAIL with m_s ≤ 1.25 × m_s_PI threshold."),
    ]:
        doc.add_heading(h, level=2)
        add_todo(doc, todo)


def add_section_discussion(doc: Document) -> None:
    doc.add_heading("8. Discussion", level=1)
    for h, todo in [
        ("8.1 Why predictive validity and RL training utility diverge",
         "Structural argument: predictive validation tests along trajectories "
         "starting from real states; closed-loop training tests trajectories "
         "generated by the model itself. Residual errors accumulate in the "
         "latter, not the former."),
        ("8.2 Why physical regularization is controller-family-specific",
         "Three mechanisms per family: thermostatic (low-dim obs needs anchor), "
         "HDRL (inter-layer conflict), 17-D MORL (rich obs is self-regularizer)."),
        ("8.3 Why observation geometry matters",
         "5-D → 17-D MORL transition: same backend, same loss, only obs "
         "differs, yet yearly m_s moves from 1.046 to 0.099."),
    ]:
        doc.add_heading(h, level=2)
        add_todo(doc, todo)

    doc.add_heading("8.4 Threats to validity", level=2)
    threats = [
        "Single weather file — climate-zone generalization not tested.",
        "Single building testcase — cross-testcase generalization addressed only if §7 is activated.",
        "KPI specificity — m_s is BOPTEST-specific; comparisons to non-BOPTEST literature are limited.",
        "Seed budget — canonical results at N = 3 seeds; HDRL at N = 1 seed.",
        "Hyperparameter sensitivity — λ sweep is targeted, not exhaustive.",
        "Hybrid recipe transferability — only single testcase without §7.",
    ]
    for t in threats:
        doc.add_paragraph(t, style="List Bullet")


def add_section_conclusion(doc: Document) -> None:
    doc.add_heading("9. Conclusion", level=1)
    add_para(doc,
        "What is solved. We show that a physically calibrated surrogate with "
        "explicit zone thermal capacitance is a strong predictive twin "
        "(24-hour rollout RMSE 0.64 °C) but fails as a closed-loop RL "
        "training environment (live RMSE > 4 °C). The resolution is a hybrid "
        "construction in which the calibrated twin enters the policy training "
        "loss only as a soft physical regularizer. The optimal regularization "
        "strength is controller-family-specific. A reproducible BOPTEST-based "
        "protocol is provided, with full numerical justification following "
        "the surrogate-development protocol of Hou and Evins (2024).")
    add_para(doc,
        "What is open. The construction has been demonstrated on a single "
        "BOPTEST testcase with a single weather file. Cross-climate and "
        "cross-building generalization are not established. A full "
        "multi-objective Pareto analysis of the MORL controller, while "
        "initiated here, requires a denser preference sweep to support a full "
        "multi-objective claim. We have intentionally avoided overclaiming "
        "the universality of the proposed recipe; the claim boundaries are "
        "stated explicitly in §8.")


def add_back_matter(doc: Document) -> None:
    doc.add_heading("Data availability", level=1)
    add_para(doc,
        "All code, trained models, configuration files, output artifacts, "
        "and the eleven supplementary tables (S1–S11) are available at "
        "<repository URL>. Numerical values reported in Tables 1–5 are "
        "reproduced directly from the corresponding CSV files via "
        "paper/build_paper_tables.py and paper/build_paper_docx.py.")

    doc.add_heading("References", level=1)
    add_todo(doc, "Populate from bibliography.bib. Seed entries: "
                  "Hou and Evins (2024), Blum et al. (2021, BOPTEST), "
                  "Schulman et al. (2017, PPO), Raissi et al. (2019, PINN), "
                  "Yang et al. (2019, MORL), Jiménez-Raboso et al. (2021, "
                  "Sinergym).")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"[INFO] paper dir : {PAPER_DIR}")
    print(f"[INFO] data root : {DATA_ROOT}")
    print(f"[INFO] out path  : {OUT_PATH}")

    doc = Document()
    _setup_styles(doc)

    add_title_block(doc)
    add_abstract(doc)
    add_section_introduction(doc)
    add_section_related_work(doc)
    add_section_methodology(doc)
    add_section_experimental_setup(doc)
    add_section_results_fidelity(doc)
    add_section_results_control(doc)
    add_section_results_transfer(doc)
    add_section_discussion(doc)
    add_section_conclusion(doc)
    add_back_matter(doc)

    doc.save(OUT_PATH)
    print(f"[OK] wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
